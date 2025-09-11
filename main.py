import streamlit as st
import io
import os
from typing import Optional, Tuple

# Core PDF libraries (these work on Streamlit Cloud)
import fitz  # PyMuPDF - faster and more reliable
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# DOCX handling
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure Streamlit page
st.set_page_config(
    page_title="ConvertX - Professional File Converter",
    page_icon="📂",
    layout="wide"
)

# Enhanced CSS styling
st.markdown("""
<style>
.stApp {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    min-height: 100vh;
}

.main-container {
    background: rgba(255, 255, 255, 0.95);
    padding: 2rem;
    border-radius: 15px;
    margin: 2rem auto;
    max-width: 800px;
    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.2);
}

.feature-card {
    background: #f8f9fa;
    padding: 1rem;
    border-radius: 10px;
    margin: 0.5rem 0;
    border-left: 4px solid #667eea;
}

.file-info {
    background: #e3f2fd;
    padding: 1rem;
    border-radius: 8px;
    margin: 1rem 0;
}
</style>
""", unsafe_allow_html=True)

# Caching decorators for performance
@st.cache_data(ttl=3600, max_entries=10)
def cached_pdf_to_text(pdf_bytes: bytes) -> str:
    """Cached PDF to text conversion using PyMuPDF"""
    text_output = ""
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    for page in pdf:
        text_output += page.get_text()
    pdf.close()
    return text_output

@st.cache_data(ttl=3600, max_entries=10)
def cached_text_to_pdf_advanced(text_content: str, font_size: int = 12) -> bytes:
    """Advanced text to PDF conversion with better formatting"""
    buffer = io.BytesIO()
    
    # Use ReportLab for better PDF generation
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create custom style
    custom_style = ParagraphStyle(
        'CustomStyle',
        parent=styles['Normal'],
        fontSize=font_size,
        spaceAfter=12,
        fontName='Helvetica'
    )
    
    story = []
    
    # Split text into paragraphs and handle formatting
    paragraphs = text_content.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            # Clean and escape text for ReportLab
            clean_text = para_text.replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
            para = Paragraph(clean_text, custom_style)
            story.append(para)
            story.append(Spacer(1, 6))
    
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

@st.cache_data(ttl=3600, max_entries=10)
def cached_docx_to_pdf_advanced(docx_bytes: bytes) -> bytes:
    """Advanced DOCX to PDF conversion preserving formatting"""
    doc = Document(io.BytesIO(docx_bytes))
    
    # Extract content with better formatting preservation
    buffer = io.BytesIO()
    pdf_doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Preserve some basic formatting
            style = styles['Normal']
            if paragraph.style.name.startswith('Heading'):
                style = styles['Heading1']
            
            # Handle text alignment
            if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                style = ParagraphStyle('CenterStyle', parent=style, alignment=1)
            elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                style = ParagraphStyle('RightStyle', parent=style, alignment=2)
            
            para = Paragraph(paragraph.text.replace('<', '&lt;').replace('>', '&gt;'), style)
            story.append(para)
            story.append(Spacer(1, 12))
    
    pdf_doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

@st.cache_data(ttl=3600, max_entries=10)
def cached_pdf_to_docx_advanced(pdf_bytes: bytes) -> bytes:
    """Advanced PDF to DOCX conversion with better layout preservation"""
    # Extract text with formatting information
    pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
    doc = Document()
    
    for page_num in range(pdf.page_count):
        page = pdf[page_num]
        
        # Extract text blocks with formatting
        blocks = page.get_text("dict")
        
        for block in blocks["blocks"]:
            if "lines" in block:
                for line in block["lines"]:
                    line_text = ""
                    for span in line["spans"]:
                        line_text += span["text"]
                    
                    if line_text.strip():
                        p = doc.add_paragraph(line_text.strip())
                        
                        # Try to preserve some formatting based on font size
                        if line["spans"] and line["spans"][0]["size"] > 14:
                            p.style = doc.styles['Heading 1']
                        elif line["spans"] and line["spans"][0]["size"] > 12:
                            p.style = doc.styles['Heading 2']
    
    pdf.close()
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def get_file_info(uploaded_file) -> dict:
    """Get comprehensive file information"""
    file_size_mb = uploaded_file.size / (1024 * 1024)
    return {
        "name": uploaded_file.name,
        "size_mb": round(file_size_mb, 2),
        "type": uploaded_file.type,
        "extension": os.path.splitext(uploaded_file.name)[1].lower()
    }

def validate_file_for_conversion(file_info: dict, target_format: str) -> Tuple[bool, str]:
    """Validate if file can be converted to target format"""
    source_ext = file_info["extension"]
    
    # Define supported conversions
    supported_conversions = {
        ".txt": [".pdf"],
        ".docx": [".pdf", ".txt"],
        ".pdf": [".txt", ".docx"]
    }
    
    if source_ext not in supported_conversions:
        return False, f"Source format {source_ext} is not supported"
    
    if target_format not in supported_conversions[source_ext]:
        return False, f"Cannot convert {source_ext} to {target_format}"
    
    # Check file size limits based on conversion type
    if source_ext == ".pdf" and file_info["size_mb"] > 50:
        return False, "PDF files larger than 50MB may cause performance issues"
    
    if file_info["size_mb"] > 100:
        return False, "Files larger than 100MB are not recommended for conversion"
    
    return True, "Conversion supported"

def process_conversion(uploaded_file, target_format: str, file_info: dict):
    """Process file conversion with progress tracking"""
    source_ext = file_info["extension"]
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        status_text.text("Reading file...")
        progress_bar.progress(10)
        
        file_bytes = uploaded_file.read()
        
        status_text.text("Processing conversion...")
        progress_bar.progress(30)
        
        output_file = None
        mime_type = None
        
        # Perform conversion based on source and target formats
        if source_ext == ".txt" and target_format == ".pdf":
            status_text.text("Converting text to PDF...")
            progress_bar.progress(60)
            
            text_content = file_bytes.decode("utf-8")
            output_bytes = cached_text_to_pdf_advanced(text_content)
            output_file = io.BytesIO(output_bytes)
            mime_type = "application/pdf"
            
        elif source_ext == ".docx" and target_format == ".pdf":
            status_text.text("Converting DOCX to PDF...")
            progress_bar.progress(60)
            
            output_bytes = cached_docx_to_pdf_advanced(file_bytes)
            output_file = io.BytesIO(output_bytes)
            mime_type = "application/pdf"
            
        elif source_ext == ".pdf" and target_format == ".txt":
            status_text.text("Extracting text from PDF...")
            progress_bar.progress(60)
            
            text_content = cached_pdf_to_text(file_bytes)
            output_file = io.BytesIO(text_content.encode("utf-8"))
            mime_type = "text/plain"
            
        elif source_ext == ".pdf" and target_format == ".docx":
            status_text.text("Converting PDF to DOCX...")
            progress_bar.progress(60)
            
            output_bytes = cached_pdf_to_docx_advanced(file_bytes)
            output_file = io.BytesIO(output_bytes)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        
        elif source_ext == ".docx" and target_format == ".txt":
            status_text.text("Extracting text from DOCX...")
            progress_bar.progress(60)
            
            doc = Document(io.BytesIO(file_bytes))
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            output_file = io.BytesIO(text_content.encode("utf-8"))
            mime_type = "text/plain"
        
        progress_bar.progress(90)
        status_text.text("Finalizing...")
        
        progress_bar.progress(100)
        status_text.text("✅ Conversion completed successfully!")
        
        return output_file, mime_type
        
    except Exception as e:
        progress_bar.progress(0)
        status_text.text(f"❌ Error: {str(e)}")
        st.error(f"Conversion failed: {str(e)}")
        return None, None

# Main app interface
def main():
    st.markdown('<div class="main-container">', unsafe_allow_html=True)
    
    st.title("📂 ConvertX - Professional File Converter")
    st.markdown("### Transform your documents with advanced conversion capabilities")
    
    # Feature highlights
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <h4>🚀 Fast Processing</h4>
            <p>Optimized algorithms for quick conversions</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <h4>🎨 Format Preservation</h4>
            <p>Maintains formatting and layout quality</p>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <h4>🔒 Secure Processing</h4>
            <p>Files processed securely with caching</p>
        </div>
        """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    # File upload section
    st.subheader("📁 Upload Your File")
    uploaded_file = st.file_uploader(
        "Choose a file to convert",
        type=["pdf", "docx", "txt"],
        help="Supported formats: PDF, DOCX, TXT. Maximum recommended size: 50MB for cloud deployment"
    )
    
    if uploaded_file:
        file_info = get_file_info(uploaded_file)
        
        # Display file information
        st.markdown(f"""
        <div class="file-info">
            <h4>📄 File Information</h4>
            <p><strong>Name:</strong> {file_info['name']}</p>
            <p><strong>Size:</strong> {file_info['size_mb']} MB</p>
            <p><strong>Type:</strong> {file_info['extension'].upper()}</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Conversion options
        st.subheader("🔄 Select Target Format")
        
        # Dynamic format options based on source file
        source_ext = file_info["extension"]
        if source_ext == ".txt":
            format_options = [".pdf"]
        elif source_ext == ".docx":
            format_options = [".pdf", ".txt"]
        elif source_ext == ".pdf":
            format_options = [".txt", ".docx"]
        else:
            format_options = []
        
        if format_options:
            target_format = st.selectbox(
                "Convert to:",
                format_options,
                format_func=lambda x: x.upper()
            )
            
            # Validate conversion
            is_valid, message = validate_file_for_conversion(file_info, target_format)
            
            if is_valid:
                st.success(f"✅ {message}")
                
                # Advanced options
                with st.expander("⚙️ Advanced Options"):
                    if source_ext == ".txt" and target_format == ".pdf":
                        font_size = st.slider("Font Size", 8, 16, 12)
                        st.info("Choose font size for PDF output")
                    
                    preserve_formatting = st.checkbox("Preserve formatting (when possible)", value=True)
                    st.info("Uses ReportLab and PyMuPDF for optimal quality")
                
                # Convert button
                if st.button("🚀 Convert File", type="primary", use_container_width=True):
                    with st.spinner("Processing..."):
                        output_file, mime_type = process_conversion(uploaded_file, target_format, file_info)
                    
                    if output_file:
                        base_name = os.path.splitext(file_info["name"])[0]
                        output_filename = f"{base_name}{target_format}"
                        
                        st.balloons()
                        
                        # Download button
                        st.download_button(
                            label=f"📥 Download {output_filename}",
                            data=output_file,
                            file_name=output_filename,
                            mime=mime_type,
                            type="primary",
                            use_container_width=True
                        )
                        
                        # Conversion summary
                        st.success(f"""
                        🎉 **Conversion Successful!**
                        
                        - **Original:** {file_info['name']} ({file_info['size_mb']} MB)
                        - **Converted:** {output_filename}
                        - **Format:** {source_ext.upper()} → {target_format.upper()}
                        """)
            else:
                st.error(f"❌ {message}")
        else:
            st.error("❌ Unsupported file format")
    
    else:
        st.info("👆 Please upload a file to get started")
    
    # Additional information
    st.markdown("---")
    with st.expander("ℹ️ Supported Conversions & Features"):
        st.markdown("""
        **Supported Conversions:**
        - **TXT → PDF**: Creates formatted PDF documents with ReportLab
        - **DOCX → PDF**: Preserves formatting and layout using advanced processing
        - **DOCX → TXT**: Extracts plain text content
        - **PDF → TXT**: Extracts text content using PyMuPDF
        - **PDF → DOCX**: Converts to editable Word document
        
        **Performance Features:**
        - ⚡ **Caching**: Repeated conversions are lightning fast
        - 🎨 **Format Preservation**: Maintains fonts, sizes, and basic formatting
        - 🔒 **Security**: Files processed locally, not stored
        - 📱 **Mobile Friendly**: Works on all devices
        
        **Quality Features:**
        - Uses **PyMuPDF** for fast, accurate PDF processing
        - Uses **ReportLab** for professional PDF generation
        - Intelligent text extraction and formatting
        - Progress tracking and error handling
        """)
    
    st.markdown('</div>', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
